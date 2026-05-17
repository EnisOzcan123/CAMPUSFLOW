from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "teslim_dokumanlari"

INK = "182026"
BLUE = "315F8C"
GREEN = "2F6F5E"
GOLD = "B7791F"
MUTED = "66717C"
LIGHT = "F5F8FB"
LINE = "DDE5EC"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=LINE, size="6"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=120, start=140, bottom=120, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_run(run, size=None, bold=None, color=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph(paragraph, before=0, after=6, line=1.1, align=None):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if align is not None:
        paragraph.alignment = align


def add_text(doc, text, size=11, color=INK, bold=False, after=6, before=0, align=None):
    p = doc.add_paragraph()
    set_paragraph(p, before=before, after=after, align=align)
    r = p.add_run(text)
    set_run(r, size=size, color=color, bold=bold)
    return p


def add_heading(doc, text, level=1):
    style = f"Heading {level}"
    p = doc.add_paragraph(style=style)
    if level == 1:
        set_paragraph(p, before=14, after=7, line=1.1)
        size, color = 16, BLUE
    elif level == 2:
        set_paragraph(p, before=11, after=5, line=1.1)
        size, color = 13, BLUE
    else:
        set_paragraph(p, before=8, after=4, line=1.1)
        size, color = 12, INK
    r = p.add_run(text)
    set_run(r, size=size, color=color, bold=True)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_paragraph(p, after=4, line=1.15)
    if p.runs:
        p.runs[0].text = text
    else:
        p.add_run(text)
    for run in p.runs:
        set_run(run, size=10.5, color=INK)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    set_paragraph(p, after=4, line=1.15)
    p.add_run(text)
    for run in p.runs:
        set_run(run, size=10.5, color=INK)
    return p


def add_callout(doc, title, body, tone="blue"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = False
    table.columns[0].width = Inches(6.3)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "EEF5F2" if tone == "green" else "F1F5FA")
    set_cell_border(cell, GREEN if tone == "green" else BLUE, "8")
    set_cell_margins(cell, 160, 180, 160, 180)
    p = cell.paragraphs[0]
    set_paragraph(p, after=3)
    r = p.add_run(title)
    set_run(r, size=11, bold=True, color=GREEN if tone == "green" else BLUE)
    p2 = cell.add_paragraph()
    set_paragraph(p2, after=0, line=1.15)
    r2 = p2.add_run(body)
    set_run(r2, size=10.5, color=INK)
    doc.add_paragraph()


def add_table(doc, rows, widths):
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = False
    for i, w in enumerate(widths):
        table.columns[i].width = Inches(w)
    for i, head in enumerate(rows[0]):
        cell = table.cell(0, i)
        set_cell_shading(cell, LIGHT)
        set_cell_border(cell, LINE)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        set_paragraph(p, after=0)
        r = p.add_run(head)
        set_run(r, size=10.2, bold=True, color=INK)
    for row in rows[1:]:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cell = cells[i]
            set_cell_border(cell, LINE)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            set_paragraph(p, after=0, line=1.12)
            r = p.add_run(text)
            set_run(r, size=10, color=INK)
    doc.add_paragraph()
    return table


def setup_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1


def header_footer(doc, label):
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = label
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph(header, after=0)
    for run in header.runs:
        set_run(run, size=9, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.text = "CampusFlow | Web Programlama Dersi"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(footer, after=0)
    for run in footer.runs:
        set_run(run, size=9, color=MUTED)


def cover(doc, title, subtitle, meta, report_type):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = False
    table.columns[0].width = Inches(6.5)
    cell = table.cell(0, 0)
    set_cell_shading(cell, INK)
    set_cell_border(cell, INK)
    set_cell_margins(cell, 360, 360, 360, 360)

    p = cell.paragraphs[0]
    set_paragraph(p, after=8)
    r = p.add_run(report_type.upper())
    set_run(r, size=10, bold=True, color=GOLD)

    p2 = cell.add_paragraph()
    set_paragraph(p2, after=8)
    r2 = p2.add_run(title)
    set_run(r2, size=25, bold=True, color=WHITE)

    p3 = cell.add_paragraph()
    set_paragraph(p3, after=12, line=1.15)
    r3 = p3.add_run(subtitle)
    set_run(r3, size=12.5, color="D8E2EA")

    p4 = cell.add_paragraph()
    set_paragraph(p4, after=0)
    r4 = p4.add_run(meta)
    set_run(r4, size=10.5, bold=True, color="D8E2EA")

    doc.add_paragraph()


def markdown_to_doc(doc, md_text, skip_title=True):
    lines = md_text.splitlines()
    in_table = []
    skipped_h1 = False
    for raw in lines:
        line = raw.rstrip()
        if not line.strip():
            if in_table:
                flush_table(doc, in_table)
                in_table = []
            continue
        if line.startswith("|"):
            in_table.append(line)
            continue
        if in_table:
            flush_table(doc, in_table)
            in_table = []

        if line.startswith("# "):
            if skip_title and not skipped_h1:
                skipped_h1 = True
                continue
            add_heading(doc, line[2:].strip(), 1)
        elif line.startswith("## "):
            add_heading(doc, line[3:].strip(), 1)
        elif line.startswith("### "):
            add_heading(doc, line[4:].strip(), 2)
        elif line.startswith("- "):
            add_bullet(doc, line[2:].strip())
        elif re.match(r"^\d+\. ", line):
            add_number(doc, re.sub(r"^\d+\. ", "", line))
        elif line.startswith("> "):
            add_callout(doc, "Not", line[2:].strip(), "green")
        else:
            add_text(doc, line, size=10.8, color=INK, after=6)
    if in_table:
        flush_table(doc, in_table)


def flush_table(doc, lines):
    rows = []
    for line in lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", c or "") for c in cells):
            continue
        rows.append(cells)
    if not rows:
        return
    col_count = len(rows[0])
    if col_count == 2:
        widths = [1.7, 4.6]
    elif col_count == 3:
        widths = [1.3, 2.6, 2.4]
    else:
        widths = [6.3 / col_count] * col_count
    add_table(doc, rows, widths)


def build_system_narrative():
    doc = Document()
    setup_styles(doc)
    header_footer(doc, "CampusFlow Sistem Anlatısı")
    cover(
        doc,
        "CampusFlow Sistem Anlatısı",
        "Kampüs mekanları, etkinlikler, harita, hava durumu ve bilet akışının sade açıklaması.",
        "Hazırlanan çıktı: Arkadaşa / ekip üyesine aktarılacak sistem özeti",
        "Sistem Rehberi",
    )
    add_callout(
        doc,
        "Kısa özet",
        "CampusFlow, öğrencinin kampüste yer bulma, etkinlik takip etme, hava durumuna göre karar verme ve bilet alma ihtiyaçlarını tek sistemde toplar.",
        "green",
    )
    markdown_to_doc(doc, (ROOT / "SISTEM_ANLATISI.md").read_text(encoding="utf-8"))
    path = OUT / "CampusFlow_Sistem_Anlatisi.docx"
    doc.save(path)
    return path


def build_project_report():
    doc = Document()
    setup_styles(doc)
    header_footer(doc, "CampusFlow Proje Raporu")
    cover(
        doc,
        "CampusFlow Proje Raporu",
        "Web Programlama Dersi dönem projesi için kısa, açık ve yönergeye uygun rapor.",
        "Proje: CampusFlow | Yapay zeka kullanımı: %60 ekip geliştirmesi, %40 destek",
        "Dönem Projesi Raporu",
    )
    add_callout(
        doc,
        "Rapor kapsamı",
        "Bu rapor proje konusu, hedef kullanıcı, temel özellikler, teknolojiler, veritabanı, özgün özellikler, yapay zeka kullanımı ve katkı beyanını içerir.",
        "blue",
    )
    markdown_to_doc(doc, (ROOT / "PROJE_RAPORU.md").read_text(encoding="utf-8"))
    path = OUT / "CampusFlow_Proje_Raporu.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    OUT.mkdir(exist_ok=True)
    print(build_system_narrative())
    print(build_project_report())
